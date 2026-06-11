"""Candidate package loading. Deterministic; declares what it skips."""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

SUPPORTED = {".md", ".txt", ".html", ".htm", ".pdf", ".docx"}

URL = re.compile(r"https?://[^\s)\"'>]+")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        return "[pdf ingestion unavailable: install pypdf]"
    try:
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:
        return f"[pdf could not be parsed: {e}]"


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return "[docx ingestion unavailable: install python-docx]"
    try:
        doc = Document(str(path))
        # Paragraph text includes every run regardless of font colour or the
        # hidden attribute, so white-on-white or hidden text inside the document
        # is extracted and reaches the detector, exactly as it does for PDFs.
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:  # resumes often lay content out in tables
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(t for t in parts if t)
    except Exception as e:
        return f"[docx could not be parsed: {e}]"


def _read_text(path: Path) -> str:
    """Read a candidate file as text. PDF and DOCX are extracted here, which
    means parsing happens wherever the pipeline runs, including inside the
    no-network sandbox. Document parsers are a known exploit surface, so we keep
    parsing in the hardened pipeline and never in a browser."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    return path.read_text(encoding="utf-8", errors="replace")


@dataclass
class CandidatePackage:
    candidate_id: str
    root: Path
    files: List[Path] = field(default_factory=list)
    skipped: List[Path] = field(default_factory=list)
    external_links: List[str] = field(default_factory=list)

    def combined_text(self) -> str:
        parts = []
        for f in self.files:
            parts.append(f"=== FILE: {f.name} ===\n{_read_text(f)}")
        return "\n\n".join(parts)


def load_candidate(root: Path) -> CandidatePackage:
    pkg = CandidatePackage(candidate_id=root.name, root=root)
    for f in sorted(root.iterdir()):
        if not f.is_file():
            continue
        if f.suffix.lower() in SUPPORTED:
            pkg.files.append(f)
            pkg.external_links.extend(URL.findall(_read_text(f)))
        else:
            pkg.skipped.append(f)
    # de-dup links, preserve order
    seen = set()
    pkg.external_links = [u for u in pkg.external_links if not (u in seen or seen.add(u))]
    return pkg
